import os
import sys
import re
import time
from rich.console import Console
from rich.progress import Progress, BarColumn, TextColumn, TimeElapsedColumn, SpinnerColumn
from rich.panel import Panel
from rich.table import Table
from rich.prompt import Prompt

# Initialize Rich Console
console = Console()

def contains_chinese(text):
    """Detects if a string contains any Chinese characters."""
    return bool(re.search(r'[\u4e00-\u9fff]', text))

def runs_have_same_formatting(run1, run2):
    """Checks if two runs have identical formatting."""
    return (run1.bold == run2.bold and
            run1.italic == run2.italic and
            run1.underline == run2.underline and
            run1.font.name == run2.font.name and
            run1.font.size == run2.font.size and
            run1.font.color.rgb == run2.font.color.rgb)

def merge_runs_in_paragraph(paragraph):
    """Merges consecutive runs with identical formatting to improve translation."""
    if len(paragraph.runs) < 2:
        return
    i = len(paragraph.runs) - 1
    while i > 0:
        run_current = paragraph.runs[i]
        run_prev = paragraph.runs[i-1]
        
        # Merge if same formatting AND regular text (no drawings/images)
        if (runs_have_same_formatting(run_prev, run_current) and
            not run_current._element.xpath(".//w:drawing") and
            not run_prev._element.xpath(".//w:drawing")):
            
            run_prev.text += run_current.text
            # Delete current run
            p_element = paragraph._p
            p_element.remove(run_current._element)
        i -= 1

def smart_translate(text, target_lang, zh_to_en, en_to_zh):
    """Translates text based on detected content and target preference."""
    from deep_translator import GoogleTranslator
    if not text or not text.strip() or len(text.strip()) < 1:
        return text
    
    has_cn = contains_chinese(text)
    
    try:
        if target_lang == 'en':
            # Want English. If has Chinese, translate it.
            if has_cn:
                res = zh_to_en.translate(text)
                return res if res is not None else text
            return text
        elif target_lang in ['zh-TW', 'zh-CN']:
            # Want Chinese. If it's English, translate it.
            if not has_cn and len(text.strip()) > 3: # Avoid very short labels
                res = en_to_zh.translate(text)
                return res if res is not None else text
            return text
        else:
            # General auto-translation for other languages
            res = GoogleTranslator(source='auto', target=target_lang).translate(text)
            return res if res is not None else text
    except:
        return text

def process_translation(file_path, target_lang):
    from docx import Document
    from deep_translator import GoogleTranslator
    
    if not os.path.exists(file_path):
        console.print(f"[bold red]❌ File not found: {file_path}[/bold red]")
        return None

    # Count blocks
    # We delay counting until we load, but we want an immediate bar
    with Progress(
        SpinnerColumn(),
        TextColumn("[progress.description]{task.description}"),
        BarColumn(bar_width=None, pulse_style="blue"),
        TimeElapsedColumn(),
        console=console
    ) as progress:
        load_task = progress.add_task(f"[yellow]📂 Parsing {os.path.basename(file_path)}... (Large file, please wait)", total=None)
        doc = Document(file_path)
        progress.update(load_task, description="[yellow]🏗️ Optimizing layout structure...", completed=100, total=100)
        
        # Merging
        for para in doc.paragraphs: merge_runs_in_paragraph(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs: merge_runs_in_paragraph(para)
        
        progress.update(load_task, description="[bold green]✅ Loaded and optimized!", completed=100)

    # Pre-init translators
    zh_to_en = GoogleTranslator(source='zh-TW', target='en')
    en_to_zh = GoogleTranslator(source='en', target='zh-TW')

    # Count total blocks for translation
    total_blocks = len(doc.paragraphs) + sum(len(cell.paragraphs) for table in doc.tables for row in table.rows for cell in row.cells)
    
    # Translation Progress
    with Progress(
        SpinnerColumn(),
        TextColumn("[progress.description]{task.description}"),
        BarColumn(bar_width=None, pulse_style="blue"),
        TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
        TimeElapsedColumn(),
        console=console
    ) as progress:
        task = progress.add_task("[bold green]✨ Translating to English...", total=total_blocks)
        
        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                for run in para.runs:
                    if run.text.strip():
                        run.text = smart_translate(run.text, target_lang, zh_to_en, en_to_zh)
            progress.update(task, advance=1)
            
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            for run in para.runs:
                                if run.text.strip():
                                    run.text = smart_translate(run.text, target_lang, zh_to_en, en_to_zh)
                        progress.update(task, advance=1)

    # Save
    base_name, ext = os.path.splitext(file_path)
    save_path = f"{base_name} (English){ext}"
    
    with console.status("[bold green]💾 Saving finalized document...", spinner="aesthetic"):
        doc.save(save_path)
    
    return save_path

def main():
    while True:
        console.clear()
        console.print(Panel.fit(
            "\n[bold magenta]━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━[/bold magenta]\n"
            "[bold white]    💎 NEURO-TRANS PREMIER : ENGLISH EDITION 💎    [/bold white]\n"
            "[dim]    Medical-Grade Layout Preservation • Smart Context Mapping    [/dim]\n"
            "[bold magenta]━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━[/bold magenta]\n",
            border_style="cyan",
            title="[bold blue]SYSTEM ONLINE[/bold blue]",
            padding=(1, 4)
        ))
        
        # File/Folder Selection
        console.print("\n[bold yellow]📂 SELECT SOURCE DOCUMENT OR FOLDER FOR ENGLISH RECONSTRUCTION:[/bold yellow]")
        
        all_entries = [f for f in os.listdir('.') if not f.startswith('.')]
        folders = [f for f in all_entries if os.path.isdir(f) and f != '.venv' and f != '__pycache__']
        files = [f for f in all_entries if f.endswith('.docx') and ' (English)' not in f]
        
        items = []
        for f in sorted(folders):
            items.append({"name": f, "type": "DIR"})
        for f in sorted(files):
            items.append({"name": f, "type": "FILE"})
            
        if not items:
            console.print("[bold green]🌟 ALL ARCHIVES HAVE BEEN RECONSTRUCTED TO ENGLISH! ✨[/bold green]")
            break
            
        for i, item in enumerate(items, 1):
            icon = "📁" if item["type"] == "DIR" else "📄"
            label = f"[DIR] " if item["type"] == "DIR" else ""
            console.print(f" [bold cyan]{i}[/bold cyan]. {icon} {label}{item['name']}")
        console.print(f" [bold red]Q[/bold red]. EXIT SYSTEM")
        
        choice = Prompt.ask("\n[bold]Select ARCHIVE ID[/bold]", default="1")
        
        if choice.lower() == 'q':
            console.print("[bold red]SYSTEM OFFLINE. GOODBYE.[/bold red]")
            break
            
        if choice.isdigit() and 1 <= int(choice) <= len(items):
            target_item = items[int(choice)-1]
            target_path = target_item["name"]
            is_dir = target_item["type"] == "DIR"
        else:
            target_path = choice if os.path.exists(choice) else None
            is_dir = os.path.isdir(target_path) if target_path else False
            
        if not target_path:
            console.print("[bold red]❌ INVALID ARCHIVE IDENTIFIED.[/bold red]")
            time.sleep(2)
            continue

        start_time = time.time()
        
        if is_dir:
            # Process folder
            docx_files = [os.path.join(target_path, f) for f in os.listdir(target_path) 
                          if f.endswith('.docx') and ' (English)' not in f]
            
            if not docx_files:
                console.print(f"[bold yellow]⚠️ No valid .docx files found in folder: {target_path}[/bold yellow]")
                time.sleep(2)
                continue
                
            console.print(f"\n[bold magenta]📦 BATCH PROCESSING: {len(docx_files)} files found in {target_path}[/bold magenta]\n")
            
            for index, docx_file in enumerate(docx_files, 1):
                console.print(f"[bold blue]━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━[/bold blue]")
                console.print(f"[bold cyan]📄 [{index}/{len(docx_files)}] Processing: {os.path.basename(docx_file)}[/bold cyan]")
                process_translation(docx_file, "en")
                console.print(f"[bold green]✅ Finished: {os.path.basename(docx_file)}[/bold green]\n")
            
            result_display = f"Folder: {target_path} ({len(docx_files)} files)"
            result = True
        else:
            # Process single file
            result = process_translation(target_path, "en")
            result_display = os.path.basename(result) if result else None
        
        if result:
            end_time = time.time()
            elapsed = end_time - start_time
            console.print(Panel(
                f"\n[bold green]🏆 ARCHIVE RECONSTRUCTED SUCCESSFULLY![/bold green]\n\n"
                f"  [bold blue]OUTPUT:[/bold blue] {result_display}\n"
                f"  [bold blue]ENGINE RUNTIME:[/bold blue] {elapsed:.1f}s",
                border_style="green",
                title="✨ Final Report ✨",
                subtitle="[dim]Check the source location for the translated copies[/dim]"
            ))
            
            cont = Prompt.ask("\n[bold]WOULD YOU LIKE TO RECONSTRUCT ANOTHER ARCHIVE?[/bold] (y/n)", choices=["y", "n"], default="y")
            if cont.lower() != 'y':
                console.print("[bold red]SYSTEM OFFLINE. GOODBYE.[/bold red]")
                break

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        console.print("\n[bold red]Operation cancelled by user.[/bold red]")
    except Exception as e:
        console.print(f"\n[bold red]An error occurred: {e}[/bold red]")

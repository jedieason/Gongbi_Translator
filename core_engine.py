import os
import re
from docx import Document
from deep_translator import GoogleTranslator
from concurrent.futures import ThreadPoolExecutor

def contains_chinese(text):
    """Detects if a string contains any Chinese characters."""
    return bool(re.search(r'[\u4e00-\u9fff]', text))

def runs_have_same_formatting(run1, run2):
    """Checks if two runs have identical formatting."""
    try:
        return (run1.bold == run2.bold and
                run1.italic == run2.italic and
                run1.underline == run2.underline and
                run1.font.name == run2.font.name and
                run1.font.size == run2.font.size and
                run1.font.color.rgb == run2.font.color.rgb)
    except:
        return False

def merge_runs_in_paragraph(paragraph):
    """Merges consecutive runs with identical formatting to improve translation."""
    if len(paragraph.runs) < 2:
        return
    i = len(paragraph.runs) - 1
    while i > 0:
        run_current = paragraph.runs[i]
        run_prev = paragraph.runs[i-1]
        try:
            if (runs_have_same_formatting(run_prev, run_current) and
                not run_current._element.xpath(".//w:drawing") and
                not run_prev._element.xpath(".//w:drawing")):
                run_prev.text += run_current.text
                p_element = paragraph._p
                p_element.remove(run_current._element)
        except:
            pass
        i -= 1

def google_translate(text, target_lang='en'):
    """Translates text using Google Translate."""
    if not text or not text.strip() or len(text.strip()) < 1:
        return text
    has_cn = contains_chinese(text)
    try:
        if target_lang == 'en':
            if has_cn:
                return GoogleTranslator(source='zh-TW', target='en').translate(text)
            return text
        elif target_lang in ['zh-TW', 'zh-CN']:
            if not has_cn and len(text.strip()) > 3:
                return GoogleTranslator(source='en', target='zh-TW').translate(text)
            return text
        else:
            return GoogleTranslator(source='auto', target=target_lang).translate(text)
    except:
        return text

def process_document(file_path, output_path, progress_callback=None):
    """
    Translates a Word document (Chinese → English) using Google Translate.
    Saves the result to output_path.
    Returns (output_path, error_message).
    """
    if not os.path.exists(file_path):
        return None, "File not found."

    try:
        doc = Document(file_path)
    except Exception as e:
        return None, f"Failed to open document: {e}"

    # Step 1: Optimise run structure
    for para in doc.paragraphs:
        merge_runs_in_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    merge_runs_in_paragraph(para)

    # Step 2: Collect all runs with text
    all_runs = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text.strip():
                all_runs.append(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip():
                            all_runs.append(run)

    total_runs = len(all_runs)
    if total_runs == 0:
        return None, "No text content found in the document."

    # Step 3: Translate runs in parallel
    completed = [0]

    def translate_run(run_obj):
        run_obj.text = google_translate(run_obj.text)
        completed[0] += 1
        if progress_callback:
            progress_callback(completed[0], total_runs)

    with ThreadPoolExecutor(max_workers=10) as executor:
        executor.map(translate_run, all_runs)

    # Step 4: Save to the user-specified output path
    try:
        doc.save(output_path)
    except Exception as e:
        return None, f"Failed to save document: {e}"

    return output_path, None
